import docx
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def add_heading(doc, text, level=1):
    heading = doc.add_heading(text, level=level)
    return heading

def add_paragraph(doc, text, style='Normal'):
    p = doc.add_paragraph(text, style=style)
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    return p

def add_image_placeholder(doc, instructions):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"\n[IMAGE PLACEHOLDER]\n{instructions}\n")
    run.bold = True
    run.font.color.rgb = docx.shared.RGBColor(255, 0, 0)
    p.style = 'Normal'

def add_code_block(doc, code):
    p = doc.add_paragraph(code)
    p.style = 'Normal'
    p.paragraph_format.left_indent = Inches(0.5)
    for run in p.runs:
        run.font.name = 'Courier New'

doc = docx.Document()
style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(12)

# Title Page
doc.add_heading('Machine Maintenance & Motor Diagnostic System', 0)
add_paragraph(doc, '\nA Z-Score Based Real-Time Anomaly Detection Approach\n').alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_page_break()

# Table of Contents placeholder
doc.add_heading('CONTENTS', level=1)
add_paragraph(doc, "Abstract\nChapter I\n1.1. Introduction\n1.1.1. Project Overview\n1.1.2. Problem Statement\nChapter II\n2.1. Literature Review\nChapter III\n3.1. Methodology\n3.1.1. Mathematical Formulation & Algorithm\n3.1.2. Hardware Architecture\n3.1.3. Software Architecture\nChapter IV\n4.1. Result and discussion\n4.1.1. Analysis Performed\n4.1.2. Discussion\nChapter V\n5.1. Conclusion\n5.2. References")
doc.add_page_break()

# Abstract
add_heading(doc, 'Abstract')
abstract_text = (
    "In contemporary industrial settings, unexpected motor and machine failures culminate in exorbitant downtime and disrupted manufacturing operations. "
    "This comprehensive report details the conceptualization, design, and implementation of a real-time Machine Maintenance and Motor Diagnostic System. "
    "Engineered to preemptively pinpoint mechanical faults through intricate vibration analysis, the system integrates a multi-axis sensor array powered by an ESP32 microcontroller. "
    "By continuously streaming X, Y, and Z-axis vibration telemetry over serial communication to a high-performance Python FastAPI backend, the architecture establishes a robust data acquisition pipeline. "
    "To accurately distinguish between normal operational variations and critical mechanical anomalies, the system employs a statistical Z-Score anomaly detection algorithm based on the Empirical Rule (3-Sigma limit). "
    "This mathematically rigorous approach dynamically calculates deviations from an initially calibrated baseline of 1000 readings. "
    "The user interface, developed utilizing Streamlit and Plotly.js, provides an intuitive frontend dashboard characterized by real-time telemetry visualization, statistical distribution profiling, and instantaneous alerting mechanisms. "
    "Ultimately, this architecture offers a highly scalable, mathematically profound, and computationally efficient edge-to-cloud solution for predictive maintenance, shifting paradigms from reactive repairs to proactive machine management."
)
add_paragraph(doc, abstract_text)
doc.add_page_break()

# Chapter I
add_heading(doc, 'Chapter I', level=1)
add_heading(doc, '1.1. Introduction', level=2)
add_paragraph(doc, 
    "Predictive maintenance represents a monumental paradigm shift from traditional reactive and scheduled machine management toward proactive, data-driven operational intelligence. "
    "Historically, industries have relied on routine maintenance, replacing parts on fixed schedules regardless of their actual degradation state, or run-to-failure strategies, repairing machinery only after catastrophic breakdowns. "
    "Both methodologies are inherently flawed; the former results in wasted resources and unnecessary downtime, while the latter risks catastrophic failure, collateral damage, and unmitigated production halts. "
    "By continuously monitoring the physical state of industrial machinery, particularly through vibrational signatures, anomalies that precede catastrophic failures can be identified well in advance. "
    "This project implements a comprehensive hardware-software pipeline to monitor three-dimensional vibration patterns, leveraging the computational capabilities of modern edge devices and high-throughput backend services."
)

add_heading(doc, '1.1.1. Project Overview', level=3)
add_paragraph(doc, 
    "The structural architecture of this diagnostic system consists of a deeply integrated embedded sensor node that transmits continuous serial data to a FastAPI-driven backend framework. "
    "At the edge, an MPU6050 accelerometer coupled with an ESP32 microcontroller samples vibration data at a high frequency. "
    "This data represents the structural dynamic behavior of the monitored motor or machine. "
    "Once transmitted via standard serial protocols, the backend asynchronously processes the telemetry, maintaining a rolling state of statistical distributions. "
    "A modern, asynchronous Streamlit and client-side Plotly.js frontend visualizes this live telemetry. The visualization is specifically optimized to highlight anomalies with zero-flicker latency, ensuring operators have instantaneous feedback regarding machine health. "
    "The core diagnostic logic relies on the real-time statistical normalization—specifically, the calculation of standard Z-Scores—of incoming sensor data against a calibrated, baseline threshold established during optimal machine operation."
)

add_heading(doc, '1.1.2. Problem Statement', level=3)
add_paragraph(doc, 
    "Industrial machines inevitably experience gradual mechanical degradation over time. This wear and tear often manifests as abnormal vibrations, acoustic emissions, or thermal fluctuations long before ultimate, critical failure occurs. "
    "Traditional maintenance schedules, derived from manufacturer estimations or historical averages, fail to account for the unique operating conditions, load variances, and environmental stresses of individual machines. "
    "Consequently, maintenance is frequently performed either too early, squandering capital on viable components, or too late, precipitating catastrophic breakdowns. "
    "The primary objective of this project is to construct an automated, real-time diagnostic tool that autonomously establishes a statistical baseline of 'normal' operational vibrations. "
    "By employing robust mathematical models, specifically the 3-Sigma limits, the system dynamically flags statistically significant deviations. These deviations serve as early warning indicators, alerting operators of impending mechanical issues such as bearing wear, rotor unbalance, or structural looseness before they propagate into critical systemic failures."
)
doc.add_page_break()

# Chapter II
add_heading(doc, 'Chapter II', level=1)
add_heading(doc, '2.1. Literature Review', level=2)
add_paragraph(doc, 
    "The evolution of Machine Maintenance has transitioned through multiple distinct phases: Reactive Maintenance, Preventive Maintenance, and presently, Predictive Maintenance (PdM). "
    "Reactive maintenance, often colloquially termed 'firefighting,' operates on the principle of repairing assets post-failure. While simple to conceptualize, its economic ramifications in high-throughput industrial sectors are severely detrimental due to unplanned downtime and compromised safety. "
    "Preventive Maintenance mitigates this by scheduling periodic interventions based on temporal or usage metrics. However, studies indicate that up to 30% of preventive maintenance is superfluous, addressing components that have significant remaining useful life (RUL). "
    "\n\nIn recent decades, Predictive Maintenance (PdM) has emerged as the industry gold standard, predicated on condition-based monitoring. "
    "A multitude of academic and industrial studies have conclusively demonstrated that vibration analysis is among the most sensitive and reliable indicators of electro-mechanical health. "
    "Vibrational signatures inherently encapsulate the harmonic frequencies of rotating machinery; deviations from these harmonic baselines are direct corollaries of physical anomalies such as misalignments, bearing faults, or shaft bows. "
    "\n\nWhile contemporary literature heavily emphasizes complex Machine Learning (ML) and Deep Learning (DL) models—such as Convolutional Neural Networks (CNNs) for spectrographic analysis or Long Short-Term Memory (LSTM) networks for time-series forecasting—these approaches often suffer from high computational overhead, necessitating substantial data lakes for training and lacking deterministic interpretability. "
    "Conversely, Statistical Process Control (SPC) methods, notably the application of Z-Scores and the Empirical Rule (3-Sigma limits), remain ubiquitous within industrial standards. "
    "The Z-Score methodology is characterized by its remarkable computational efficiency, enabling microsecond latency on edge computing nodes. "
    "Furthermore, it offers absolute interpretability, operating effectively on normally distributed vibrational data without the prerequisite of exhaustive failure datasets, which are notoriously difficult to acquire in highly reliable manufacturing environments. "
    "\n\nThe architectural paradigm of Edge-to-Cloud integrations has further revolutionized PdM. By decentralizing data acquisition to microcontrollers (such as the ESP32) equipped with micro-electromechanical systems (MEMS) sensors (like the MPU6050), and centralizing analytical computation within scalable web frameworks (like FastAPI), systems can achieve unprecedented scalability."
)
doc.add_page_break()

# Chapter III
add_heading(doc, 'Chapter III', level=1)
add_heading(doc, '3.1. Methodology', level=2)
add_paragraph(doc, 
    "The methodological framework of this project is architected around a continuous, highly synchronized loop comprising three distinct phases: Data Acquisition, Baseline Calibration, and Real-Time Statistical Monitoring. "
    "This deterministic approach guarantees that the diagnostic system adapts to the specific operating signature of the attached machinery, circumventing the need for generalized, potentially inaccurate, factory models."
)

add_heading(doc, '3.1.1. Mathematical Formulation & Algorithm', level=3)
add_paragraph(doc, 
    "The foundational logic governing the anomaly detection engine is heavily rooted in standard statistical probability, specifically the properties of normal (Gaussian) distributions. "
    "Vibrational data in properly functioning, continuously rotating machinery generally exhibits a normal distribution centered around a specific mean amplitude. "
    "\n\nStep 1: Calibration (Baseline Establishment)\n"
    "During an authenticated 'healthy' operational state, the system transitions into a calibration phase. It records an empirical dataset consisting of N = 1000 absolute vibration readings across three spatial dimensions (X, Y, and Z axes). "
    "Absolute values are strictly utilized to prevent negative and positive oscillatory data points from mathematically canceling each other out, which would erroneously yield a mean approaching zero. "
    "For each axis, the baseline Statistical Mean (μ) and Standard Deviation (σ) are computed using the following standard statistical formulas:"
)
add_image_placeholder(doc, "INSERT DIAGRAM HERE: Diagram showing the equations for Mean (μ) and Standard Deviation (σ). You can generate these using LaTeX or screenshot them from the markdown presentation.")

add_paragraph(doc, 
    "Step 2: Real-Time Z-Score Computation\n"
    "Upon successful calibration, the system transitions into active monitoring. For every subsequent, real-time incoming reading (|x|), a Z-Score is computed. "
    "The Z-Score is a dimensionless quantity that represents the precise number of standard deviations a given data point resides away from the established mean. "
    "This normalization is critical as it standardizes the variance across the X, Y, and Z axes, regardless of their intrinsic amplitude differences."
)
add_image_placeholder(doc, "INSERT DIAGRAM HERE: Diagram showing the Z-Score formula: Z = (|x| - μ) / σ.")

add_paragraph(doc, 
    "Step 3: Anomaly Threshold (Decision Rule)\n"
    "In accordance with the Empirical Rule associated with normal distributions, approximately 99.73% of all healthy data points will naturally fall within three standard deviations (±3σ) of the mean. "
    "Therefore, the system adopts a rigid detection rule: If the absolute value of the computed Z-Score exceeds 3 (|Z| > 3), the data point is classified as statistically highly improbable under normal operating conditions. "
    "Consequently, the system instantaneously flags this event as a critical Anomaly. The safe operating boundaries are thus mathematically defined as [μ - 3σ, μ + 3σ]."
)
add_image_placeholder(doc, "INSERT GRAPH HERE: Insert the 'gaussian_deviation_bands.png' here. This beautifully illustrates the normal distribution and the 3-sigma anomaly thresholds.")

add_heading(doc, '3.1.2. Hardware Architecture', level=3)
add_paragraph(doc, 
    "The hardware subsystem is engineered for robust data acquisition directly at the mechanical edge. The primary processing unit is the ESP32 microcontroller, selected for its high clock speed, dual-core architecture, and reliable serial communication interfaces. "
    "Interfaced with the ESP32 via the I2C protocol is the Adafruit MPU6050, a highly sensitive 6-axis MotionTracking device that combines a 3-axis gyroscope and a 3-axis accelerometer. "
    "\n\nFor the scope of this project, the accelerometer is configured to a range of ±8g, providing sufficient headroom to capture violent mechanical shocks while maintaining high resolution for subtle vibrations. "
    "A digital low-pass filter with a bandwidth of 21Hz is applied directly at the hardware level to attenuate high-frequency electrical noise and mechanical resonance that fall outside the spectrum of typical machine degradation signatures. "
    "The firmware dictates a precise sampling loop, reading the sensor registers and transmitting JSON-formatted payloads over the UART serial interface at a strict interval of 20 milliseconds (50 Hz sampling rate)."
)
add_image_placeholder(doc, "INSERT HARDWARE IMAGE HERE: A clear, well-lit photograph showing the ESP32 microcontroller wired to the MPU6050 sensor. Add a caption: 'Figure: Embedded Edge Device Assembly showcasing ESP32 and MPU6050'.")

add_paragraph(doc, "Below is the core C++ firmware logic running on the ESP32 edge device, demonstrating the precise timing loop and data transmission protocol:")
cpp_code = """#include <Arduino.h>
#include <Adafruit_MPU6050.h>
#include <Adafruit_Sensor.h>
#include <Wire.h>

Adafruit_MPU6050 mpu;
unsigned long lastSendTime = 0;

void setup() {
  Serial.begin(115200);
  if (!mpu.begin()) {
    Serial.println("{\\"error\\": \\"MPU6050 not found\\"}");
    while (1);
  }
  mpu.setAccelerometerRange(MPU6050_RANGE_8_G);
  mpu.setFilterBandwidth(MPU6050_BAND_21_HZ);
}

void loop() {
  if (millis() - lastSendTime >= 20) {
    lastSendTime = millis(); 
    sensors_event_t a, g, temp;
    mpu.getEvent(&a, &g, &temp);

    char strX[16], strY[16], strZ[16];
    sprintf(strX, "%.4f", a.acceleration.x);
    sprintf(strY, "%.4f", a.acceleration.y);
    sprintf(strZ, "%.4f", a.acceleration.z); 

    char jsonPayload[128];
    sprintf(jsonPayload, "{\\"x\\": %s, \\"y\\": %s, \\"z\\": %s}", strX, strY, strZ);
    Serial.println(jsonPayload);
  }
}"""
add_code_block(doc, cpp_code)

add_heading(doc, '3.1.3. Software Architecture', level=3)
add_paragraph(doc, 
    "The software pipeline bridges the edge hardware with a sophisticated data processing and visualization stack. "
    "The backend is constructed using FastAPI, a modern, fast (high-performance) web framework for building APIs with Python. "
    "FastAPI was selected over traditional frameworks like Flask or Django due to its native asynchronous support, which is paramount for concurrently handling continuous serial I/O streams and serving HTTP API requests without blocking. "
    "\n\nA dedicated background daemon thread (`serial_worker`) maintains a persistent connection to the serial port, parsing incoming JSON payloads. "
    "During the calibration phase, this thread writes the absolute vibrational values to a localized CSV database (`calibration_1000.csv`). "
    "During standard operation, it updates global state variables that the API endpoints interrogate to compute the real-time Z-Scores. "
    "The Z-Score history is maintained in a rolling buffer to facilitate time-series visualization on the frontend."
)

add_paragraph(doc, "The following Python snippet highlights the FastAPI backend logic responsible for ingesting serial data and managing the calibration sequence:")
py_code = """def serial_worker():
    global latest_reading, is_calibrating, calibration_count
    try:
        ser = serial.Serial(SERIAL_PORT, BAUD_RATE, timeout=1)
    except Exception:
        return

    while True:
        try:
            if ser.in_waiting > 0:
                line = ser.readline().decode('utf-8').strip()
                if line.startswith('{') and line.endswith('}'):
                    data = json.loads(line)
                    latest_reading = {
                        "x": abs(data['x']),
                        "y": abs(data['y']),
                        "z": abs(data['z'])
                    }
                    if is_calibrating and calibration_count < 1000:
                        with open(CSV_FILE, mode='a', newline='') as file:
                            csv.writer(file).writerow([abs(data['x']), abs(data['y']), abs(data['z'])])
                        calibration_count += 1
                        if calibration_count >= 1000:
                            is_calibrating = False
        except Exception:
            time.sleep(0.01)"""
add_code_block(doc, py_code)

add_paragraph(doc, 
    "The frontend presentation layer is built utilizing Streamlit. To overcome Streamlit's native limitations regarding rapid, flicker-free updates (which typically necessitate full server-side reruns), custom HTML/JS components leveraging Plotly.js were injected. "
    "This hybrid approach allows the browser client to independently poll the FastAPI backend and dynamically update the Plotly canvases in-place (`Plotly.react()`). "
    "This ensures a completely fluid, high-framerate oscilloscope-style visualization that accurately reflects the high-speed telemetry generated by the ESP32."
)
doc.add_page_break()

# Chapter IV
add_heading(doc, 'Chapter IV', level=1)
add_heading(doc, '4.1. Result and discussion', level=2)
add_heading(doc, '4.1.1. Analysis Performed', level=3)
add_paragraph(doc, 
    "The implemented system underwent rigorous testing paradigms to validate its efficacy in identifying structural anomalies. "
    "1. Calibration & Distribution Analysis: The primary analytical step involved plotting the frequency distributions of the 1000 data points collected during the calibration phase. "
    "Utilizing the Seaborn and Matplotlib libraries, probability density functions were overlaid atop the empirical histograms for each spatial axis. "
    "Visual and statistical analysis emphatically confirmed that the absolute vibration data closely approximated a normal (Gaussian) distribution. "
    "This crucial finding validated the foundational premise of employing the Z-Score method. The mean and ±3σ anomaly boundaries were accurately rendered, establishing the deterministic thresholds for subsequent monitoring."
)
add_image_placeholder(doc, "INSERT GRAPH HERE: Take a screenshot of the Streamlit application showing the 'Step 2: Distribution' output. This should show the 3 blue, green, and red histograms for the X, Y, and Z axes generated by Action 3.")

add_paragraph(doc, 
    "2. Real-Time Telemetry Simulation: The integration of the asynchronous FastAPI backend and the client-side Plotly.js frontend was evaluated for latency and rendering stability. "
    "The system successfully maintained a high-frequency, zero-flicker live chart. The implementation of a rolling visual window (configured to 100 data points) successfully mimicked the behavior of industrial-grade diagnostic oscilloscopes, providing operators with a clear, immediate view of the vibrational waveform."
)
add_image_placeholder(doc, "INSERT UI SCREENSHOT HERE: Capture a screenshot of the 'Step 4: Live Anomaly Detection' dashboard showing the Z-Score chart and the X, Y, Z axis cards indicating normal operation (green).")

add_paragraph(doc, 
    "3. Anomaly Introduction and Detection: To simulate mechanical failure, the sensor assembly was subjected to deliberate, abnormal physical shocks and induced harmonic vibrations that exceeded the baseline profile. "
    "In every instance, the calculated live Z-Score immediately spiked, breaching the predefined ±3 threshold. "
    "The system demonstrated sub-second responsiveness, instantly triggering the UI's critical alert mechanisms. The interface dynamically updated to feature a flashing red anomaly banner, and the specific axis cards exhibiting the deviation transitioned to an alarmed state, clearly delineating the vector of the structural fault."
)
add_image_placeholder(doc, "INSERT UI SCREENSHOT HERE: Capture a screenshot of the dashboard while it is detecting an Anomaly. The banner should be red ('ANOMALY DETECTED') and the Z-score chart should show peaks exceeding the dashed lines.")

add_heading(doc, '4.1.2. Discussion', level=3)
add_paragraph(doc, 
    "The analytical results substantiate the robustness of the chosen methodological framework. "
    "The explicit architectural decision to utilize absolute values (|x|) during both calibration and real-time computation proved essential. "
    "In raw vibrational telemetry, mechanical oscillations produce alternating positive and negative values. Computing a standard mean on raw data often results in an arithmetic cancellation, yielding an artificially low mean approaching zero and subsequently suppressing the calculated variance. "
    "By tracking absolute magnitudes, the system accurately characterizes the true energy envelope of the vibration. "
    "\n\nFurthermore, the system exhibited exceptional responsiveness and low latency. "
    "A significant engineering challenge in web-based dashboards is the rendering bottleneck associated with high-frequency data streams. "
    "By strategically offloading the chart rendering responsibilities to the client's browser utilizing Plotly.react, rather than forcing server-side Streamlit component reruns, the system achieved a paradigm-shifting reduction in visual latency. "
    "This UI optimization is critical for real-time fault detection. "
    "\n\nFinally, the purely statistical Z-Score approach demonstrated profound advantages over contemporary ML-based alternatives within the context of edge deployment. "
    "Neural network approaches often necessitate extensive, computationally expensive training phases and massive datasets containing labeled 'failure' states, which are rare and difficult to synthesize accurately. "
    "The Z-Score algorithm, conversely, requires zero model training time. It is highly suitable for rapid deployment and scaling across heterogeneous machine types. "
    "Each individual machine requires only a brief, 20-second (1000-tick) calibration phase to establish its unique operational signature, after which the system is fully autonomous and armed for anomaly detection."
)
doc.add_page_break()

# Chapter V
add_heading(doc, 'Chapter V', level=1)
add_heading(doc, '5.1. Conclusion', level=2)
add_paragraph(doc, 
    "This project successfully conceptualized, engineered, and implemented a highly scalable, mathematically rigorous Machine Maintenance and Motor Diagnostic System. "
    "By seamlessly bridging embedded edge hardware (ESP32 and MPU6050) with a robust mathematical anomaly detection engine based on definitive Z-Score analysis, and exposing the analytics through a high-performance web dashboard, the architecture provides a comprehensive, enterprise-grade solution for predictive maintenance. "
    "The lightweight statistical approach proved emphatically effective at differentiating between acceptable operational mechanical noise and critical, out-of-band structural deviations. "
    "The system's ability to adapt dynamically to distinct machine profiles via a rapid calibration phase highlights its immense practical applicability in diverse industrial environments. "
    "Ultimately, this project demonstrates that complex industrial monitoring challenges can frequently be resolved with elegant statistical mathematics and optimized software engineering, precluding the necessity for opaque and computationally demanding machine learning models."
)

add_heading(doc, '5.2. References', level=2)
add_paragraph(doc, 
    "1. Montgomery, D. C. (2019). Introduction to Statistical Quality Control. John Wiley & Sons.\n"
    "2. Mobley, R. K. (2002). An Introduction to Predictive Maintenance. Butterworth-Heinemann.\n"
    "3. Tiangolo, S. (2023). FastAPI Documentation. Retrieved from https://fastapi.tiangolo.com/\n"
    "4. Plotly.js Scientific Graphing Library. Retrieved from https://plotly.com/javascript/\n"
    "5. Espressif Systems. ESP32 Technical Reference Manual. Retrieved from https://www.espressif.com/\n"
    "6. Streamlit Documentation. Retrieved from https://docs.streamlit.io/"
)

# Save the document
output_path = "c:/DEV/ML/projects/minor_project/Final_Project_Report.docx"
doc.save(output_path)
print(f"Successfully generated {output_path}")
